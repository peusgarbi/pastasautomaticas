from src.surgeon_signature import generate_surgeon_signature
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt
from docx import Document
from pathlib import Path


def generate_discharge_prescription(
    patient_name: str,
    patient_age: int,
    patient_gender: str,
    patient_address: str,
    surgeon_name: str,
    surgeon_crm: str,
    prescription_text: str,
    output_path: Path,
    surgery_date: str,
) -> Path:
    """
    Gera uma receita de alta em formato DOCX.

    Args:
        patient_name: Nome do paciente
        patient_age: Idade do paciente
        patient_gender: Sexo (M/F)
        patient_address: Endereço completo
        surgeon_name: Nome do cirurgião
        surgeon_crm: CRM do cirurgião
        prescription_text: Conteúdo da receita
        output_path: Caminho do DOCX de saída

    Returns:
        Path do arquivo gerado
    """

    document = Document()

    #
    # CONFIGURAÇÃO DA PÁGINA
    #

    section = document.sections[0]

    # A5 retrato
    section.page_width = Cm(14.8)
    section.page_height = Cm(21)

    # Margens padrão do hospital
    section.top_margin = Cm(4.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    #
    # FONTE PADRÃO
    #

    style = document.styles["Normal"]

    style.font.name = "Calibri"
    style.font.size = Pt(10)

    #
    # CABEÇALHO
    #

    header = document.add_paragraph()

    header.paragraph_format.space_before = Pt(0)
    header.paragraph_format.space_after = Pt(0)

    run = header.add_run(
        f"Nome: {patient_name}    {patient_age} anos    {patient_gender}"
    )
    run.bold = True

    address = document.add_paragraph()
    run = address.add_run(f"Endereço: {patient_address}")
    run.bold = True

    address.paragraph_format.space_before = Pt(0)
    address.paragraph_format.space_after = Pt(0)

    document.add_paragraph("")

    #
    # CORPO DA RECEITA
    #

    body = document.add_paragraph()

    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(0)
    body.paragraph_format.line_spacing = 1

    body.add_run(prescription_text)

    #
    # DATA
    #

    document.add_paragraph()

    city_date = document.add_paragraph()
    run = city_date.add_run(f"São José do Rio Preto, {surgery_date}")
    run.bold = True

    city_date.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    #
    # ASSINATURA
    #

    document.add_paragraph()
    signature = document.add_paragraph()
    signature.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    signature.paragraph_format.left_indent = Cm(4.5)
    run = signature.add_run(generate_surgeon_signature(surgeon_name))
    run.bold = True

    #
    # SALVAR
    #

    output_dir = Path(output_path)

    output_dir.mkdir(
        parents=True,
        exist_ok=True,
    )

    filename = f"{patient_name} - RECEITA ALTA.docx"

    file_path = output_dir / filename

    document.save(file_path)

    return file_path
