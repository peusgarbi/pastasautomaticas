from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt


def generate_companion_declaration(
    patient_name: str,
    patient_age: int,
    patient_gender: str,
    patient_address: str,
    surgeon_name: str,
    surgeon_crm: str,
    output_dir: Path,
    surgery_date: str,
) -> Path:

    document = Document()

    #
    # CONFIGURAÇÃO DA PÁGINA
    #

    section = document.sections[0]

    section.page_width = Cm(14.8)
    section.page_height = Cm(21)

    section.top_margin = Cm(4.0)
    section.bottom_margin = Cm(2.5)

    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    section.footer_distance = Cm(2.0)

    #
    # ESTILO
    #

    style = document.styles["Normal"]

    style.font.name = "Calibri"
    style.font.size = Pt(10)

    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1

    #
    # CABEÇALHO
    #

    header = document.add_paragraph()

    header.paragraph_format.space_before = Pt(0)
    header.paragraph_format.space_after = Pt(0)

    run = header.add_run(
        f"Nome: {patient_name}    {patient_age} anos    {patient_gender}"
    )
    run.bold = True

    address = document.add_paragraph()
    run = address.add_run(f"Endereço: {patient_address}")
    run.bold = True

    address.paragraph_format.space_before = Pt(0)
    address.paragraph_format.space_after = Pt(0)

    document.add_paragraph("")

    #
    # TÍTULO
    #

    title = document.add_paragraph()

    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = title.add_run("DECLARAÇÃO MÉDICA")
    run.bold = True
    run.underline = True
    run.font.size = Pt(14)

    document.add_paragraph("")

    #
    # CORPO
    #

    body = document.add_paragraph()
    body.add_run(
        "Declaro para os devidos fins que o(a) Sr(a). "
        "____________________________________________ "
        f"acompanhou o(a) paciente {patient_name} "
        "durante procedimento cirúrgico realizado "
        "nesta instituição nesta data.\n\n"
        "Necessita de afastamento de suas atividades "
        "por _____ (__________) dias, "
        "a contar da presente data."
    )

    #
    # CID
    #

    cid = document.add_paragraph("")
    run = cid.add_run("CID-10: Z76.3")
    run.bold = True

    #
    # DATA
    #

    document.add_paragraph("")

    city_date = document.add_paragraph()
    run = city_date.add_run(f"São José do Rio Preto, {surgery_date}")
    run.bold = True
    city_date.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    #
    # ASSINATURA
    #

    footer = section.footer

    footer_paragraph = footer.paragraphs[0]

    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    footer_paragraph.paragraph_format.left_indent = Cm(5.5)

    footer_run = footer_paragraph.add_run(
        f"{surgeon_name}\nMédico Otorrinolaringologista\nCRM {surgeon_crm}"
    )

    footer_run.bold = True

    #
    # SALVAR
    #

    output_dir = Path(output_dir)

    output_dir.mkdir(
        parents=True,
        exist_ok=True,
    )

    file_path = output_dir / f"{patient_name} - DECLARACAO ACOMPANHANTE.docx"

    document.save(file_path)

    return file_path
